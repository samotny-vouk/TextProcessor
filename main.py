import sys
from PyQt5.QtWidgets import QApplication, QMainWindow, QTextEdit, QAction, QMenu, QFileDialog, QInputDialog, \
    QColorDialog, QFontDialog, QMessageBox
from PyQt5.QtGui import QFont, QTextCharFormat, QPixmap, QColor
from PyQt5.QtCore import Qt, QLocale
# import pypdf2
from docx import Document
import zipfile


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()

        self.doc = Document()
        self.default_style = self.doc.styles['Normal']

        self.fileMenu = self.menuBar().addMenu("&Файл")
        self.openAction = QAction("&Открыть", self)
        self.openAction.triggered.connect(self.openFile)
        self.fileMenu.addAction(self.openAction)

        self.saveAction = QAction("&Сохранить", self)
        self.saveAction.triggered.connect(self.saveFile)
        self.fileMenu.addAction(self.saveAction)

        self.saveAsAction = QAction("Сохранить как...", self)
        self.saveAsAction.triggered.connect(self.saveFileAs)
        self.fileMenu.addAction(self.saveAsAction)

        self.exportMenu = self.fileMenu.addMenu("Экспорт")

        self.exportPdfAction = QAction("PDF", self)
        self.exportPdfAction.triggered.connect(self.exportPdf)
        self.exportMenu.addAction(self.exportPdfAction)

        self.exportTxtAction = QAction("TXT", self)
        self.exportTxtAction.triggered.connect(self.exportTxt)
        self.exportMenu.addAction(self.exportTxtAction)

        self.exportZipAction = QAction("ZIP", self)
        self.exportZipAction.triggered.connect(self.exportZip)
        self.exportMenu.addAction(self.exportZipAction)

        self.fileMenu.addSeparator()

        self.exitAction = QAction("&Выход", self)
        self.exitAction.triggered.connect(self.close)
        self.fileMenu.addAction(self.exitAction)

        self.editMenu = self.menuBar().addMenu("&Правка")
        self.undoAction = QAction("&Отменить", self)
        self.undoAction.triggered.connect(self.undo)
        self.editMenu.addAction(self.undoAction)

        self.redoAction = QAction("&Повторить", self)
        self.redoAction.triggered.connect(self.redo)
        self.editMenu.addAction(self.redoAction)

        self.formatMenu = self.menuBar().addMenu("&Формат")

        self.insertImageAction = QAction("Вставить изображение", self)
        self.insertImageAction.triggered.connect(self.insertImage)
        self.formatMenu.addAction(self.insertImageAction)

        self.insertLinkAction = QAction("Вставить ссылку", self)
        self.insertLinkAction.triggered.connect(self.insertLink)
        self.formatMenu.addAction(self.insertLinkAction)

        self.boldAction = QAction("&Жирный", self)
        self.boldAction.triggered.connect(self.setBold)
        self.formatMenu.addAction(self.boldAction)

        self.italicAction = QAction("&Курсив", self)
        self.italicAction.triggered.connect(self.setItalic)
        self.formatMenu.addAction(self.italicAction)

        self.underlineAction = QAction("&Подчеркивание", self)
        self.underlineAction.triggered.connect(self.setUnderline)
        self.formatMenu.addAction(self.underlineAction)

        self.fontAction = QAction("&Шрифт", self)
        self.fontAction.triggered.connect(self.changeFont)
        self.formatMenu.addAction(self.fontAction)

        self.fontSizeAction = QAction("&Размер шрифта", self)
        self.fontSizeAction.triggered.connect(self.changeFontSize)
        self.formatMenu.addAction(self.fontSizeAction)

        self.textColorAction = QAction("&Цвет текста", self)
        self.textColorAction.triggered.connect(self.changeTextColor)
        self.formatMenu.addAction(self.textColorAction)

        self.indentAction = QAction("&Отступ", self)
        self.indentAction.triggered.connect(self.applyIndent)
        self.formatMenu.addAction(self.indentAction)

        self.outdentAction = QAction("&Убрать отступ", self)
        self.outdentAction.triggered.connect(self.applyOutdent)
        self.formatMenu.addAction(self.outdentAction)

        self.lineSpacingAction = QAction("&Интервал между строками", self)
        self.lineSpacingAction.triggered.connect(self.changeLineSpacing)
        self.formatMenu.addAction(self.lineSpacingAction)

        self.textEdit = QTextEdit(self)
        self.setCentralWidget(self.textEdit)

        self.setStyleSheet("QMainWindow { background-color: #f0f0f0; }"
                          "QTextEdit { font-family: Arial; font-size: 12pt; }")

    def insertImage(self):
        options = QFileDialog.Options()
        filePath, _ = QFileDialog.getOpenFileName(self, "Выбрать изображение", "",
                                                  "Images (*.png *.jpg *.jpeg *.bmp *.gif)", options=options)
        if filePath:
            self.textEdit.insertHtml(f'<img src="{filePath}" alt="Image" width="300"/>')

    def insertLink(self):
        cursor = self.textEdit.textCursor()
        selected_text = cursor.selectedText()

        if not selected_text:
            return

        link, ok = QInputDialog.getText(self, 'Вставить ссылку', 'Введите URL:')
        if ok and link:
            char_format = cursor.charFormat()
            char_format.setForeground(Qt.blue)
            char_format.setFontUnderline(True)

            cursor.mergeCharFormat(char_format)
            cursor.insertHtml(f'<a href="{link}">{selected_text}</a>')

    def saveFileAs(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        fileName, _ = QFileDialog.getSaveFileName(self, "Сохранить как...", "",
                                                  "Документ Word (*.docx);;Текстовый файл (*.txt);;Все файлы (*)",
                                                  options=options)
        if fileName:
            self.currentFile = fileName
            self.saveDocument(fileName)

    def saveDocument(self, fileName):
        text = self.textEdit.toPlainText()
        self.doc.add_paragraph(text)
        self.doc.save(fileName)

    def exportPdf(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        fileName, _ = QFileDialog.getSaveFileName(self, "Экспорт в PDF", "",
                                                  "PDF (*.pdf);;Все файлы (*)",
                                                  options=options)
        if fileName:
            docx_file = fileName.replace('.pdf', '.docx')
            self.doc.save(docx_file)
            self.convertDocxToPdf(docx_file, fileName)

    def exportTxt(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        fileName, _ = QFileDialog.getSaveFileName(self, "Экспорт в TXT", "",
                                                  "Текстовый файл (*.txt);;Все файлы (*)",
                                                  options=options)
        if fileName:
            with open(fileName, 'w', encoding='utf-8') as f:
                f.write(self.textEdit.toPlainText())

    def exportZip(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        fileName, _ = QFileDialog.getSaveFileName(self, "Экспорт в ZIP", "",
                                                  "ZIP (*.zip);;Все файлы (*)",
                                                  options=options)
        if fileName:
            with zipfile.ZipFile(fileName, 'w', zipfile.ZIP_DEFLATED) as zipf:
                zipf.write(self.currentFile, arcname=self.currentFile.split('/')[-1])

    def convertDocxToPdf(self, docx_file, pdf_file):
        """Конвертирует DOCX в PDF."""
        try:
            from docx2pdf import convert  # Импортируйте модуль docx2pdf
            convert(docx_file, pdf_file)  # Конвертируйте DOCX в PDF
        except ImportError:
            QMessageBox.warning(self, "Ошибка",
                                "Не установлен модуль docx2pdf. Установите его с помощью 'pip install docx2pdf'")
        except AssertionError as e:
            QMessageBox.warning(self, "Ошибка", f"Ошибка конвертации: {e}")

    def openFile(self):
        fileName = QFileDialog.getOpenFileName(self, "Открыть файл", "", "Текстовые файлы (*.txt);;Все файлы (*)")[0]
        if fileName:
            with open(fileName, 'r', encoding='utf-8') as file:
                self.textEdit.setText(file.read())

    def saveFile(self):
        fileName = QFileDialog.getSaveFileName(self, "Сохранить файл", "", "Текстовые файлы (*.txt);;Все файлы (*)")[0]
        if fileName:
            with open(fileName, 'w', encoding='utf-8') as file:
                file.write(self.textEdit.toPlainText())

    def undo(self):
        self.textEdit.undo()

    def redo(self):
        self.textEdit.redo()

    def setBold(self):
        cursor = self.textEdit.textCursor()
        fmt = cursor.charFormat()
        if self.textEdit.currentFont().weight() == QFont.Normal:
            fmt.setFontWeight(QFont.Black)
        elif self.textEdit.currentFont().weight() == QFont.Black:
            fmt.setFontWeight(QFont.Normal)
        else:
            fmt.setFontWeight(QFont.Normal)
        cursor.mergeCharFormat(fmt)

    def setItalic(self):
        cursor = self.textEdit.textCursor()
        fmt = cursor.charFormat()
        fmt.setFontItalic(not self.textEdit.currentFont().italic())
        cursor.mergeCharFormat(fmt)

    def setUnderline(self):
        cursor = self.textEdit.textCursor()
        fmt = cursor.charFormat()
        currentUnderline = fmt.underlineStyle()
        if currentUnderline == QTextCharFormat.UnderlineStyle.NoUnderline:
            fmt.setUnderlineStyle(QTextCharFormat.UnderlineStyle.SingleUnderline)
        elif currentUnderline == QTextCharFormat.UnderlineStyle.SingleUnderline:
            fmt.setUnderlineStyle(QTextCharFormat.UnderlineStyle.DashUnderline)
        elif currentUnderline == QTextCharFormat.UnderlineStyle.DashUnderline:
            fmt.setUnderlineStyle(QTextCharFormat.UnderlineStyle.DotLine)
        elif currentUnderline == QTextCharFormat.UnderlineStyle.DotLine:
            fmt.setUnderlineStyle(QTextCharFormat.UnderlineStyle.DashDotLine)
        elif currentUnderline == QTextCharFormat.UnderlineStyle.DashDotLine:
            fmt.setUnderlineStyle(QTextCharFormat.UnderlineStyle.DashDotDotLine)
        elif currentUnderline == QTextCharFormat.UnderlineStyle.DashDotDotLine:
            fmt.setUnderlineStyle(QTextCharFormat.UnderlineStyle.WaveUnderline)
        else:
            fmt.setUnderlineStyle(QTextCharFormat.UnderlineStyle.NoUnderline)

        cursor.mergeCharFormat(fmt)
        self.textEdit.setTextCursor(cursor)

    def changeFontSize(self):
        fontSize, ok = QInputDialog.getInt(self, "Размер шрифта", "Введите размер:", 10, 1, 100)
        if ok:
            cursor = self.textEdit.textCursor()
            fmt = cursor.charFormat()
            font = QFont(fmt.font())
            font.setPointSize(fontSize)
            fmt.setFont(font)
            cursor.mergeCharFormat(fmt)
            self.textEdit.setTextCursor(cursor)

    def changeFont(self):
        font, ok = QFontDialog.getFont()
        if ok:
            cursor = self.textEdit.textCursor()
            fmt = cursor.charFormat()
            fmt.setFont(font)
            cursor.setCharFormat(fmt)
            self.textEdit.setTextCursor(cursor)

    def changeTextColor(self):
        color = QColorDialog.getColor(self.textEdit.textColor())
        if color.isValid():
            self.textEdit.setTextColor(color)

    def applyIndent(self):
        cursor = self.textEdit.textCursor()
        fmt = QTextCharFormat()
        fmt.BlockIndent(cursor.blockFormat().indent() + 1)
        cursor.mergeCharFormat(fmt)

    def applyOutdent(self):
        self.applyFormat(QTextCharFormat.BlockIndent, self.textEdit.textCursor().blockFormat().indent() - 1)

    def changeLineSpacing(self):
        spacing, ok = QInputDialog.getDouble(self, "Интервал между строками", "Введите интервал:", 1.0, 0.1, 3.0, 2)
        if ok:
            self.textEdit.setLineSpacing(spacing)


if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())